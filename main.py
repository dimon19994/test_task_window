import sys

import pandas as pd
from PyQt5.QtCore import QFileInfo
from PyQt5.QtWidgets import (
    QApplication,
    QCheckBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QInputDialog,
    QLabel,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QTableWidgetItem,
)
from docx.api import Document
from openpyxl import load_workbook

from mainwindow import Ui_MainWindow


class CsvDialog(QDialog):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setWindowTitle("Select CSV details")
        self._gui_init()

    def _gui_init(self):
        layout = QFormLayout()
        self.setLayout(layout)

        self.delimiter = QLineEdit(self)
        self.is_headers = QCheckBox(self)

        layout.addRow(QLabel("Delimiter: "), self.delimiter)
        layout.addRow(QLabel("First row is headers: "), self.is_headers)

        button_box = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self
        )
        layout.addRow(button_box)

        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)

    def get_inputs(self):
        return self.delimiter.text(), self.is_headers.isChecked()


class HeadersDialog(QDialog):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setWindowTitle("Headers")
        self._gui_init()

    def _gui_init(self):
        layout = QFormLayout()
        self.setLayout(layout)

        self.is_headers = QCheckBox(self)
        layout.addRow(QLabel("First row is headers: "), self.is_headers)

        button_box = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self
        )
        layout.addRow(button_box)

        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)

    def get_inputs(self):
        return self.is_headers.isChecked()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self._display_default_table()

        self.ui.tableWidget.horizontalHeader().sectionDoubleClicked.connect(
            lambda index: self._change_header(index, "horizontalHeaderItem", "column")
        )
        self.ui.tableWidget.verticalHeader().sectionDoubleClicked.connect(
            lambda index: self._change_header(index, "verticalHeaderItem", "row")
        )

        self.ui.actionOpen_file.triggered.connect(self._open_file)
        self.ui.actionSave_file.triggered.connect(self._save_file)
        self.ui.actionAdd_row.triggered.connect(self._add_row)
        self.ui.actionAdd_column.triggered.connect(self._add_column)
        self.ui.actionDelete_row.triggered.connect(self._remove_row)
        self.ui.actionDelete_column.triggered.connect(self._remove_column)

    def _add_row(self):
        row_position = self.ui.tableWidget.rowCount()
        self.ui.tableWidget.insertRow(row_position)

        it = QTableWidgetItem(str(row_position + 1))
        self.ui.tableWidget.setVerticalHeaderItem(row_position, it)

    def _add_column(self):
        column_position = self.ui.tableWidget.columnCount()
        self.ui.tableWidget.insertColumn(column_position)

        it = QTableWidgetItem(f"Column {column_position + 1}")
        self.ui.tableWidget.setHorizontalHeaderItem(column_position, it)

    def _change_header(self, index, method, text):
        header_item = getattr(self.ui.tableWidget, method)
        old_header = header_item(index).text()
        new_header, ok = QInputDialog.getText(
            self,
            f"Change header label for {text} %d" % index,
            "Header:",
            QLineEdit.Normal,
            old_header,
        )

        if ok:
            header_item(index).setText(new_header)

    def _display_default_table(self):
        self._set_table_parameters(1, 1, ("Column 1",))

    def _load_csv(self, path, delimiter, is_headers):
        try:
            df = pd.read_csv(
                path, delimiter=delimiter, header=0 if is_headers else None
            )
        except FileNotFoundError:
            self._display_message_box(
                QMessageBox.Critical,
                "File not found",
            )
            return
        except pd.errors.EmptyDataError:
            self._display_message_box(
                QMessageBox.Critical,
                "No data",
            )
            return
        except pd.errors.ParserError:
            self._display_message_box(
                QMessageBox.Critical,
                "Parse error",
            )
            return
        except Exception:
            self._display_message_box(
                QMessageBox.Critical,
                "Some error",
            )
            return

        if is_headers:
            headers = tuple(map(str, df.head()))
        else:
            headers = (f"Column {i + 1}" for i in range(len(df.axes[1])))

        self._set_table_parameters(len(df.axes[0]), len(df.axes[1]), headers)

        for i, row in df.iterrows():
            for j in range(len(row)):
                self.ui.tableWidget.setItem(i, j, QTableWidgetItem(str(row[j])))

    def _load_docx(self, path, is_headers):
        try:
            document = Document(path)
        except FileNotFoundError:
            self._display_message_box(
                QMessageBox.Critical,
                "File not found",
            )
            return

        if len(document.tables) > 0:
            table = document.tables[0]
        else:
            self._display_message_box(
                QMessageBox.Critical,
                "There are no tables in the file",
            )
            return

        if is_headers:
            headers = (cell.text for cell in table.rows[0].cells)
            position = 1
        else:
            headers = (f"Column {i + 1}" for i in range(len(table.rows)))
            position = 0

        self._set_table_parameters(
            len(table.rows) - position, len(table.columns), headers
        )

        for i, row in enumerate(table.rows[position:]):
            for j, cell in enumerate(row.cells):
                self.ui.tableWidget.setItem(i, j, QTableWidgetItem(str(cell.text)))

    def _load_xlsx(self, path, is_headers):
        try:
            workbook = load_workbook(path)
        except FileNotFoundError:
            self._display_message_box(
                QMessageBox.Critical,
                "File not found",
            )
            return
        sheet = workbook.active

        list_values = list(sheet.values)

        if is_headers:
            headers = tuple(map(str, list_values[0]))
            position = 1
        else:
            headers = (f"Column {i + 1}" for i in range(sheet.max_column))
            position = 0

        self._set_table_parameters(sheet.max_row - position, sheet.max_column, headers)

        for i, row in enumerate(list_values[position:]):
            for j, value in enumerate(row):
                self.ui.tableWidget.setItem(i, j, QTableWidgetItem(str(value)))

    def _open_file(self):
        path, file_type = QFileDialog.getOpenFileName(
            self,
            "Open table",
            "",
            "CSV Files (*.csv);;Excel Files (*.xlsx);;Word Files (*.doc *.docx)",
        )

        if path:
            if file_type == "CSV Files (*.csv)":
                dialog = CsvDialog()
                if dialog.exec():
                    self._load_csv(path, *dialog.get_inputs())
            elif file_type == "Excel Files (*.xlsx)":
                dialog = HeadersDialog()
                if dialog.exec():
                    self._load_xlsx(path, dialog.get_inputs())
            elif file_type == "Word Files (*.doc *.docx)":
                dialog = HeadersDialog()
                if dialog.exec():
                    self._load_docx(path, dialog.get_inputs())

    def _remove_row(self):
        if self.ui.tableWidget.rowCount() == 1:
            self._display_message_box(
                QMessageBox.Critical,
                "At least one row must exist",
            )
            return

        current_row = self.ui.tableWidget.currentRow()
        self.ui.tableWidget.removeRow(current_row)

    def _remove_column(self):
        if self.ui.tableWidget.columnCount() == 1:
            self._display_message_box(
                QMessageBox.Critical,
                "At least one column must exist",
            )
            return

        current_column = self.ui.tableWidget.currentColumn()
        self.ui.tableWidget.removeColumn(current_column)

    def _save_file(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Save table", "", "Word Files (*.docx)"
        )
        document = Document()

        row_count = self.ui.tableWidget.rowCount()
        column_count = self.ui.tableWidget.columnCount()

        table = document.add_table(row_count + 1, column_count)
        table.style = "Table Grid"

        for column in range(column_count):
            header = self.ui.tableWidget.horizontalHeaderItem(column)
            table.cell(0, column).text = header.text()

        for row in range(row_count):
            for column in range(column_count):
                data = self.ui.tableWidget.item(row, column)
                if data and data.text():
                    table.cell(row + 1, column).text = data.text()
                else:
                    table.cell(row + 1, column).text = ""

        if QFileInfo(path).suffix() != "docx":
            path += ".docx"

        document.save(path)

    def _set_table_parameters(self, row_count, col_count, headers):
        self.ui.tableWidget.setRowCount(row_count)
        self.ui.tableWidget.setColumnCount(col_count)

        self.ui.tableWidget.setHorizontalHeaderLabels(headers)
        self.ui.tableWidget.setVerticalHeaderLabels(
            (str(i + 1) for i in range(row_count))
        )

    @staticmethod
    def _display_message_box(icon, info, text="Error", title="Error"):
        msg = QMessageBox()
        msg.setIcon(icon)
        msg.setText(text)
        msg.setInformativeText(info)
        msg.setWindowTitle(title)
        msg.exec()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())
